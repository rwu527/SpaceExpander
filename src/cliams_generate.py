import datetime
import os
from pydoc import doc
import random
import re
import time
from reportlab.lib.pagesizes import letter
from PIL import Image, ImageChops, ImageEnhance
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Pt, Inches, RGBColor 
from PIL import Image
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from src.concrete_content import subscript_to_int

# Constants for layout (you can adjust these as needed)
MARGIN_TOP = 20
MARGIN_LEFT = 20
MARGIN_RIGHT = 20
LINE_HEIGHT = 12
PAGE_WIDTH = 8.5 * 72  # 8.5 inches in points
PAGE_BOTTOM = 11 * 72  # 11 inches in points
PAGE_WIDTH, PAGE_HEIGHT = letter


def change_font(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, level=None):
    """Add a paragraph with specified text and set font to Times New Roman, size 11."""
    if level is not None:
        paragraph = doc.add_heading(text, level=level)
    else:
        paragraph = doc.add_paragraph(text)
    
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)  
    
    paragraph.alignment = alignment

    paragraph.paragraph_format.space_before = Pt(0) 
    paragraph.paragraph_format.space_after = Pt(4)  
    # paragraph.paragraph_format.line_spacing = Pt(18)  

    return paragraph


# Function to create a Word document
def create_word_document(docx_path, output_folder):
    """Create a Word document with claims and images from the formula folder."""
    doc = Document()

    # Add title
    title = change_font(doc, 'Claims', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    # Keep spacing consistent before and after the title
    title.paragraph_format.space_after = Inches(0.2)
    title.paragraph_format.space_before = Inches(0.2)

    # Add first claim
    change_font(doc, "1. A compound of formula (I):")
    # Set formula folder path
    formula_folder = os.path.join(output_folder, 'formula')
    
    if os.path.exists(formula_folder):
        image_files = [f for f in os.listdir(formula_folder) if f.lower().endswith('.png')]

        if image_files:
            image_path = os.path.join(formula_folder, image_files[0])

            try:
                # Create a borderless table
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                cell = table.cell(0, 0)

                # Remove table borders
                tbl_pr = table._element.tblPr
                tbl_borders = OxmlElement('w:tblBorders')
                for border_tag in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_tag}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    tbl_borders.append(border)
                tbl_pr.append(tbl_borders)

                # Insert image
                with open(image_path, 'rb') as f:
                    img_run = cell.paragraphs[0].add_run()
                    img_run.add_picture(f, width=Inches(2.5))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Place (I) at the bottom right corner of the image
                number_run = cell.paragraphs[0].add_run(" (I)")
                number_run.font.size = Pt(10)
                number_run.font.name = "Times New Roman"
                number_run.baseline = -1  # Slight adjustment to make it sink a bit

                # Align the content inside the table to the bottom
                cell_element = cell._element
                v_align = OxmlElement("w:vAlign")
                v_align.set(qn("w:val"), "bottom")
                cell_element.append(v_align)

            except Exception as e:
                print(f"[ERROR] Image insertion failed: {str(e)}")
        else:
            print("[Warning] No valid images found in the formula folder.")
    else:
        print("[Error] Formula folder does not exist.")
        
    # Add "wherein:" section
    change_font(doc, "wherein:")

    # Get all subfolder names
    subfolders = [
        f for f in os.listdir(output_folder)
        if os.path.isdir(os.path.join(output_folder, f)) and f != 'formula'
    ]

    # To record processed R groups and avoid duplication
    processed_r_groups = set()
    all_valid_parts = []

    # Step 1: Extract all valid parts
    for folder_name in subfolders:
        # If the folder name contains "_"
        if "_" in folder_name:
            parts = folder_name.split("_") 
        else:
            parts = [folder_name] 
        for part in parts:
            if part.startswith("R") and part not in processed_r_groups:
                all_valid_parts.append(part)
                processed_r_groups.add(part)
            elif not part.startswith("R"):
                all_valid_parts.append(part)

    # Sort the parts for orderly presentation
    sorted_parts = sorted(
        all_valid_parts,
        key=lambda x: int(re.search(r"\d+", x).group()) if x.startswith("R") and re.search(r"\d+", x) else float('inf')
    )

    # Initialize y_position for the first folder description
    y_position = MARGIN_TOP

    # Add each part to the document
    for i, part in enumerate(sorted_parts):
        y_position = add_folder_description(
            doc,
            part,
            os.path.join(output_folder, part),
            y_position,
            is_last_folder=(i == len(sorted_parts) - 1)
        )

    # Save the Word document
    doc.save(docx_path)
    remove_blank_paragraphs(docx_path)
    

def remove_blank_paragraphs(docx_path):
    doc = Document(docx_path)

    # Iterate through paragraphs and remove empty ones
    for para in doc.paragraphs:
        if not para.text.strip():
            p = para._element  # Access the underlying XML element
            p.getparent().remove(p)

    # Save the modified document
    doc.save(docx_path)
    

def process_nest_folder(doc, part, folder_path, y_position):
    R_nest_description = ""

    # Get all subfolders in the given folder path
    subfolders = [f for f in os.listdir(folder_path) if os.path.isdir(os.path.join(folder_path, f))]

    nest_folder_path = None
    # Search for the _nest subfolder corresponding to the part
    for subfolder in subfolders:
        if part in subfolder and "_nest" in subfolder:
            nest_folder_path = os.path.join(folder_path, subfolder)
            break                    

    if nest_folder_path and os.path.isdir(nest_folder_path):
        # Get all mcs_with_r_ images in the nest folder
        mcs_images = [f for f in os.listdir(nest_folder_path) if f.startswith("mcs_with_r_") and f.lower().endswith(".png")]

        all_images = mcs_images       
        image_paths = []
        if all_images:
            for index, image_file in enumerate(all_images):
                image_path = os.path.join(nest_folder_path, image_file)
                image_paths.append(image_path)

        # Also look for 0_mcs_with_r_ images in the nest folder
        form_images = [f for f in os.listdir(nest_folder_path) if f.startswith("0_mcs_with_r_") and f.lower().endswith(".png")]

        form_paths = []
        if form_images:
            for index, image_file in enumerate(form_images):
                form_path = os.path.join(nest_folder_path, image_file)
                form_paths.append(form_path)


        # Check if the rule_description.txt exists and process it
        txt_file_path = os.path.join(nest_folder_path, "mcs_with_r.txt")
        if os.path.exists(txt_file_path):
            with open(txt_file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                nest_chain_mcs = []
                    
                for line in lines:
                    if line.startswith("Nest chain MCS:"):
                        content = line.split("Nest chain MCS:")[1].strip()
                        nest_chain_mcs.extend(content.split(", "))

                unique_mcs = list(set(nest_chain_mcs))
                unique_mcs.sort()  
                unique_mcs = ["-" + mcs for mcs in unique_mcs]
                expanded_mcs = []
                for mcs in unique_mcs:
                    mcs_expanded = mcs 
                    expanded_mcs = []
                    possible_expansions = []
                    
                    if 'CH₂' in mcs:
                        possible_expansions.append("remove_CH₂")
                        possible_expansions.append("expand_CH₂")                    
                        possible_expansions.append("replace_CH₂_with_CO")  # Add the new expansion for CH₂ -> CO
                    if not mcs.startswith('CH₂'):
                        possible_expansions.append("add_CH₂_at_start")

                    # Process each expansion method, even if some conditions are not met
                    if "remove_CH₂" in possible_expansions:
                        mcs_expanded = mcs
                        mcs_expanded = mcs_expanded.replace('CH₂', '', 1)
                        expanded_mcs.append(mcs_expanded)

                    if "expand_CH₂" in possible_expansions:
                        mcs_expanded = mcs
                        mcs_expanded = mcs_expanded.replace('CH₂', '(CH₂)₂', 1)
                        expanded_mcs.append(mcs_expanded)

                    if "add_CH₂_at_start" in possible_expansions:
                        mcs_expanded = mcs
                        dash_index = mcs_expanded.find('-')
                        if dash_index != -1:
                            mcs_expanded = mcs_expanded[:dash_index + 1] + 'CH₂' + mcs_expanded[dash_index + 1:]
                        else:
                            mcs_expanded = 'CH₂' + mcs_expanded
                        expanded_mcs.append(mcs_expanded)

                    # New expansion: Replace CH₂ with CO
                    if "replace_CH₂_with_CO" in possible_expansions:
                        mcs_expanded = mcs
                        mcs_expanded = mcs_expanded.replace('CH₂', 'CO')  # Replace all occurrences of CH₂ with CO
                        expanded_mcs.append(mcs_expanded)

                    # Further replacements to standardize the mcs_expanded
                    mcs_expanded = mcs_expanded.replace('-=', '=')
                    mcs_expanded = mcs_expanded.replace('-≡', '≡')
                    mcs_expanded = mcs_expanded.replace('()', '')
                    mcs_expanded = compress_ch2_sequence(mcs_expanded)
                    mcs_expanded = mcs_expanded.replace('()', '')

                    # Add the original mcs as well
                    expanded_mcs.append(mcs)

                    # Remove duplicates
                    expanded_mcs = list(dict.fromkeys(expanded_mcs))  

                R_nest_description = ", ".join(expanded_mcs[:-1]) + " or " + expanded_mcs[-1] + ";"
        return y_position, nest_folder_path, True, R_nest_description, image_paths, form_paths

    return y_position, nest_folder_path, False, R_nest_description, [], []


def compress_ch2_sequence(text):
    """
    Replace sequences like 'CH₂CH₂CH₂' with '(CH₂)₃'
    """
    subscript_map = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
    
    def repl(match):
        count = len(match.group(0)) // 3  
        subscript = str(count).translate(subscript_map)
        return f"(CH₂){subscript}"

    return re.sub(r'(CH₂){2,}', repl, text)



def generate_images(doc, image_paths):
    """
    Function to insert images into a Word document while maintaining the intended format.
    """
    table = doc.add_table(rows=0, cols=3)  # Create a table with 3 columns
    table.alignment = 1  # Center alignment

    row = None
    for index, image_path in enumerate(image_paths):
        if index % 3 == 0:
            row = table.add_row().cells  # Add a new row every 3 images

        # Open and process the image
        img = Image.open(image_path).convert("L")  
        img = ImageEnhance.Contrast(img).enhance(2.0)  
        img = img.crop(img.getbbox())

        img_width, img_height = img.size
        if img_width >= img_height:
            scaled_width = 1.2  # Scaled width in inches
            scaled_height = img_height * (scaled_width / img_width)
        else:
            scaled_height = 1.2
            scaled_width = img_width * (scaled_height / img_height)

        img.save("temp_img.png")  # Save processed image
        row[index % 3].paragraphs[0].add_run().add_picture("temp_img.png", width=Inches(scaled_width))

        # Add separator text (",", "or", ";")
        if index == len(image_paths) - 2:
            row[index % 3].paragraphs[0].add_run(" or")
        elif index == len(image_paths) - 1:
            row[index % 3].paragraphs[0].add_run(";")
        else:
            row[index % 3].paragraphs[0].add_run(",")

    os.remove("temp_img.png")  # Clean up temporary image file


def generate_images_2(doc, image_paths):
    """
    Function to insert images into a Word document while maintaining the intended format.
    This version ensures that images are aligned in a three-column table, and the correct
    separator (",") is placed after each image.
    """
    table = doc.add_table(rows=0, cols=3)  # Create a table with 3 columns
    table.alignment = 1  # Center alignment

    row = None
    for index, image_path in enumerate(image_paths):
        if index % 3 == 0:
            row = table.add_row().cells  # Add a new row every 3 images

        # Open and process the image
        img = Image.open(image_path).convert("L")  
        img = ImageEnhance.Contrast(img).enhance(2.0)  
        img = img.crop(img.getbbox())

        img_width, img_height = img.size
        if img_width >= img_height:
            scaled_width = 1.2  # Scaled width in inches
            scaled_height = img_height * (scaled_width / img_width)
        else:
            scaled_height = 1.2
            scaled_width = img_width * (scaled_height / img_height)

        img.save("temp_img.png")  # Save processed image
        row[index % 3].paragraphs[0].add_run().add_picture("temp_img.png", width=Inches(scaled_width))

        # Add separator text (",")
        row[index % 3].paragraphs[0].add_run(",")

    os.remove("temp_img.png")  # Clean up temporary image file


def generate_form_images(doc, image_paths):
    """
    Function to insert images into a Word document while maintaining the intended format.
    This version ensures that images are aligned in a three-column table, and the correct
    separator (",") is placed after each image.
    """
    table = doc.add_table(rows=0, cols=3)  # Create a table with 3 columns
    table.alignment = 1  # Center alignment

    row = None
    for index, image_path in enumerate(image_paths):
        if index % 3 == 0:
            row = table.add_row().cells  # Add a new row every 3 images

        # Open and process the image
        img = Image.open(image_path).convert("L")  
        img = ImageEnhance.Contrast(img).enhance(2.0)  
        img = img.crop(img.getbbox())

        img_width, img_height = img.size
        if img_width >= img_height:
            scaled_width = 1.2  # Scaled width in inches
            scaled_height = img_height * (scaled_width / img_width)
        else:
            scaled_height = 1.2
            scaled_width = img_width * (scaled_height / img_height)

        img.save("temp_img.png")  # Save processed image
        row[index % 3].paragraphs[0].add_run().add_picture("temp_img.png", width=Inches(scaled_width))

        # Add separator text (",")
        row[index % 3].paragraphs[0].add_run(",")

    os.remove("temp_img.png")  # Clean up temporary image file


seed_initialized = False  
random_choice = None  


def process_rule_description(part, folder_path):
    """Process the rule description file and return a combined description."""
    global seed_initialized, random_choice
    rule_description_path = os.path.join(os.path.dirname(folder_path), "rule_description.txt")
    combined_description = "" 
    forms_description = ""

    # List of fallback items to choose from when items_sorted_all is empty
    fallback_items = [
        "an unsubstituted C₆-C₈ heterocyclic ring containing two oxygen atoms and one sulfur atom",
        "a substituted C₅-C₆ cycloalkyl ring optionally containing one nitrogen atom",
        "a substituted C₅-C₆ cycloalkyl ring optionally containing one oxygen atom",
        "a substituted C₅-C₇ cycloalkyl ring",
        "a substituted C₄-C₅ heteroaryl ring optionally containing 1-2 heteroatoms independently selected from N, O, or S",
        "a substituted C₅-C₇ cycloalkyl ring substituted with 1-2 fluoro groups",
        "an unsubstituted C₆ heteroaryl ring substituted with one methoxy group",
        "a substituted C₅-C₆ cycloalkyl ring substituted with a combination of methyl, chloro, and nitro groups",
        "a substituted fused bicyclic aryl system comprising a benzene ring and a pyridine ring"
    ]

    # If the seed has not been initialized, initialize it once
    if not seed_initialized:
        random.seed(time.time())  # Use the current timestamp as the seed to ensure a different seed each time the program runs
        random_choice = random.choice(fallback_items)  # Choose a random item and save it
        seed_initialized = True  # Set to initialized

    if os.path.exists(rule_description_path):
        with open(rule_description_path, "r", encoding="utf-8") as file:
            lines = file.readlines()
            descriptions = set()
            ring_descriptions = set()
            in_folder_section = False
            other_rs = set()

            for line in lines:
                # Check if the content after "Folder:" contains "_", if it does, skip it
                if line.startswith("Folder:") and part in line:
                    folder_info = line.strip().split(":")[1].strip() if ":" in line else ""
                    # Only process folder names that do not contain "_", e.g., "Rn"
                    if "_" in folder_info:  
                        continue  # Skip this line

                    # If it doesn't contain "_", continue processing the folder
                    in_folder_section = True
                    other_rs.update(folder_info.split("_"))
                    continue

                if in_folder_section:
                    if line.startswith("Folder:"):
                        break  # Exit the section if a new folder starts
                    if line.strip() == "":
                        continue  # Skip empty lines
                    
                    if "No matching description found" in line:
                        continue
                    
                    clean_line = re.sub(r".*?\.smiles:", "", line).strip() 
                    if "forms" in clean_line:
                        ring_descriptions.add(clean_line.replace("forms ", ""))
                    elif clean_line:
                        descriptions.add(clean_line)

            sorted_descriptions = sorted(descriptions)
            if ring_descriptions:
                sorted_forms = sorted(ring_descriptions)

                # Step 1: Split "forms" description into parts and store as a list of tuples (main_part, r_group)
                forms_parts = []
                for form in sorted_forms:
                    parts = form.split(" with ")
                    if len(parts) == 2:
                        main_part, r_group = parts
                        forms_parts.append((main_part.strip(), r_group.strip()))
                    else:
                        forms_parts.append((parts[0].strip(), None))

                grouped_forms = {}
                for main_part, r_group in forms_parts:
                    if r_group:
                        if r_group not in grouped_forms:
                            grouped_forms[r_group] = []
                        grouped_forms[r_group].append(main_part)
                    else:
                        grouped_forms[None] = grouped_forms.get(None, []) + [main_part]

                final_forms = []
                forms_found = False
                collected_form_strs = []  # Store all form_str that matches "unsubstituted or substituted ring"

                # Process the groups
                for r_group, items in grouped_forms.items():
                    if r_group:
                        items_sorted = sorted(set(items))  
                        if len(items_sorted) > 1:
                            form_str = f"an unsubstituted/substituted {', '.join(items_sorted[:-1])} or {items_sorted[-1]} with {r_group}" 
                        else:
                            form_str = f"an unsubstituted/substituted {items_sorted[0]} with {r_group}"
                    else:
                        items_sorted = sorted(set(items)) 
                        form_str = f"an unsubstituted/substituted {', '.join(items_sorted[:-1])} or {items_sorted[-1]}" 

                    if "other similar ring" in form_str:
                        collected_form_strs.append((items_sorted, r_group))
                    else:
                        final_forms.append(form_str)

                if collected_form_strs:
                    items_sorted_all = []
                    r_group_all = [] 

                    # Collect and merge items_sorted and r_group
                    for items_sorted, r_group in collected_form_strs:
                        items_sorted_all.extend(items_sorted)
                        r_group_all.append(r_group)

                    # Remove "other similar ring" from items_sorted_all
                    items_sorted_all = [item for item in items_sorted_all if item != "other similar ring"]

                    # If items_sorted_all is empty, choose a fallback item randomly
                    if not items_sorted_all:
                        items_sorted_all = [random_choice]  

                    # Sort items_sorted and r_group
                    items_sorted_all = sorted(set(items_sorted_all))
                    r_group_all = sorted(set(r_group_all))

                    # If only one item, no need for commas or "or"/"and"
                    if len(items_sorted_all) > 1:
                        items_sorted_str = f"{', '.join(items_sorted_all[:-1])} or {items_sorted_all[-1]}"
                    else:
                        items_sorted_str = f"or {items_sorted_all}"

                    if len(r_group_all) > 1:
                        r_group_str = f"{', '.join(r_group_all[:-1])} and {r_group_all[-1]}"
                    else:
                        r_group_str = r_group_all[0]

                    # Merge items_sorted with "or" and r_group with "and"
                    forms_description = f"{items_sorted_str} with {r_group_str};"
                else:
                    forms_description = ""  #
                forms_description = forms_description.replace("['", "").replace("']", "") 

                if final_forms:
                    # Apply similar formatting for combined description
                    ring_description = "forms " + ", ".join(final_forms)
                    other_rs.discard(part)
                    if other_rs:
                        ring_description += f" with {' and '.join(other_rs)}"

                    sorted_descriptions.append(ring_description)

            if sorted_descriptions:
                sorted_descriptions = [desc.replace(" ₁ to ₅ carbon atoms", " 1 to 5 carbon atoms") for desc in sorted_descriptions]
                sorted_descriptions = [desc.replace("₄- to ₇-membered ring", "4- to 7-membered ring") for desc in sorted_descriptions]
                if "C₁-C₄ aminoalkyl" in sorted_descriptions and "C₃-C₈ aminoalkyl" in sorted_descriptions:
                    sorted_descriptions = [desc for desc in sorted_descriptions if desc not in ["C₁-C₄ aminoalkyl", "C₃-C₈ aminoalkyl"]]
                    sorted_descriptions.append("C₁-C₈ aminoalkyl")

                if "C₃-C₈ cycloalkyl" in sorted_descriptions and "C₅-C₁₀ cycloalkyl" in sorted_descriptions:
                    sorted_descriptions = [desc for desc in sorted_descriptions if desc not in ["C₃-C₈ cycloalkyl", "C₅-C₁₀ cycloalkyl"]]
                    sorted_descriptions.append("C₃-C₁₀ cycloalkyl")

                if "C₁-C₄ alkylamino" in sorted_descriptions and "C₃-C₈ alkylamino" in sorted_descriptions:
                    sorted_descriptions = [desc for desc in sorted_descriptions if desc not in ["C₁-C₄ alkylamino", "C₃-C₈ alkylamino"]]
                    sorted_descriptions.append("C₁-C₈ alkylamino")               

                if len(sorted_descriptions) > 1:
                    # If the special phrase exists, remove it from the list and append it at the end
                    if "OR', R' is a halogen atom" in sorted_descriptions:
                        sorted_descriptions.remove("OR', R' is a halogen atom")
                        sorted_descriptions.append("OR', R' is a halogen atom")
                    
                    # Join the descriptions into a single string with commas
                    combined_description = ", ".join(sorted_descriptions[:-1]) + f" or {sorted_descriptions[-1]}"
                else:
                    combined_description = sorted_descriptions[0]

    else:
        combined_description = ""  # If the file doesn't exist, set description to empty string

    if combined_description.strip().lower() == "hydrogen":
        combined_description = "C1-C4 straight or branched chain alkyl or hydrogen"
    
    return combined_description, forms_description


def custom_replace(text):
    """Replace ';' with ',' and selectively replace ' or ' with ',' 
    unless it appears in predefined phrases.
    """
    text = text.replace(";", ",")
    
    # Predefined phrases where " or " should not be replaced
    protected_phrases = [
        "aromatic acyl which may have one halogen atom or an aryl as substituted",
        "C₁-C₄ alkyl partially or totally substituted with halogen atoms",
        "C₁-C₆ straight or branched chain alkyl",
        "C₃-C₈ straight or branched chain alkyl",
        "C₁-C₁₂ straight-chain or branched acyl",
        "cycloalkyl or phenyl optionally substituted with halogen",
        "cycloalkyl or phenyl optionally substituted with hydroxyl or phenyl",
        "C₁-C₄ alkoxy partially or totally substituted with halogen atoms",
        "aromatic or heteroaromatic which can be substituted with one or more halogens",
        "phenoxy substituted by one or more halogen atoms",
        "alkylene chain, straight or branched, of from 1 to 5 carbon atoms",
        "aromatic or heteroaromatic which can be substituted with one or more halogen atoms",
        "dimethyl or diethylformamide or acetamide",
        "lower alkyl which may have one halogen atom or an aryl as substituted",
        "phenyl substituted by one or two halogen atoms",
        "mono- or di(C₁-₄)alkyl-amino(C₁-₄)alkyl",
        "R', or OR', wherein R' is a halogen atom"
    ]

    
    # Use regex to find and protect these phrases
    for phrase in protected_phrases:
        protected_text = phrase.replace(" or ", "_PROTECTED_OR_")
        text = text.replace(phrase, protected_text)

    # Replace remaining " or " occurrences
    text = text.replace(" or ", ", ")

    # Restore protected phrases
    text = text.replace("_PROTECTED_OR_", " or ")
    
    return text


def output_images_from_r_folder(doc, r_folder_path):
    """
    Insert images from the given R folder into a Word document,
    ensuring a four-column layout with correct positioning of separators ("or", ";").
    """
    images = [f for f in os.listdir(r_folder_path) if f.lower().endswith(".png")]
    table = doc.add_table(rows=0, cols=4)  # Create a table with 4 columns
    table.alignment = 1  # Center alignment

    row = None
    row_scaled_heights = []
    
    for row_start_index in range(0, len(images), 4):  # Process images in groups of four
        row_images = images[row_start_index: row_start_index + 4]
        row = table.add_row().cells  # Add a new row for the images
        
        row_scaled_heights.clear()
        
        # Calculate max height for current row
        for image_file in row_images:
            image_path = os.path.join(r_folder_path, image_file)
            img = Image.open(image_path).convert("L")
            img = img.crop(img.getbbox())

            img_width, img_height = img.size
            if img_width >= img_height:
                scaled_width = 1.1  # Adjusted to approximately match 80px
                scaled_height = img_height * (scaled_width / img_width)
            else:
                scaled_height = 1.1
                scaled_width = img_width * (scaled_height / img_height)

            row_scaled_heights.append(scaled_height)

        max_scaled_height = max(row_scaled_heights)  # Get max height of the row
        
        # Insert images into the table row
        for index, image_file in enumerate(row_images):
            image_path = os.path.join(r_folder_path, image_file)
            img = Image.open(image_path).convert("L")
            img = img.crop(img.getbbox())

            img_width, img_height = img.size
            if img_width >= img_height:
                scaled_width = 1.1
                scaled_height = img_height * (scaled_width / img_width)
            else:
                scaled_height = 1.1
                scaled_width = img_width * (scaled_height / img_height)

            temp_path = "temp_img.png"
            img.save(temp_path)

            paragraph = row[index].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(temp_path, width=Inches(scaled_width))

            # Add separator ("or" / ";") after the image
            text = " or" if row_start_index + index == len(images) - 2 else ";"
            paragraph.add_run(text)

        os.remove(temp_path)  # Clean up temporary file
        

def generate_description():
    content_options = [
        "Cyclohexane, Cyclopentane, Benzene, Toluene",
        "4- to 7-membered ring, Benzene",
        "aromatic acyl which may have one halogen atom or an aryl as substituted",
        "C3-C6 cycloalkyl, -OH, -O(C1-C4 alkoxy)",
        "5- to 7-membered ring, optionally containing an O or NH"
    ]

    return random.choice(content_options)


from src.digital import convert_subscript_to_number


def process_r_folder_recursively(doc, r_folder, r_folder_path, y_position, index=0):  
    if "_" in r_folder:
        print(f"Skipping folder: {r_folder} (contains '_')")
        return y_position

    # Initialize full_output
    full_output = f"wherein {r_folder} is " if index == 0 else f"{r_folder} is "
    contains_nest_file_in_r_folder = any("nest" in file_name for file_name in os.listdir(r_folder_path))
    
    # If there are no nested files, directly output the description and stop further processing
    combined_description, forms_description = process_rule_description(r_folder, r_folder_path)
    if not contains_nest_file_in_r_folder:
        # If no nested files, directly output the description
        full_output += combined_description + ";"
        change_font(doc, full_output)  # Directly add the output to the Word document
        return  # Stop further processing when there are no nested files

    # Handle the logic for nested folders
    y_position, nest_folder_path, success, R_nest_description, image_paths, form_paths = process_nest_folder(doc, r_folder, r_folder_path, y_position)

    # Generate the description content
    full_output += combined_description
    if combined_description == "":
        full_output += f"{generate_description()}"

    if image_paths:
        full_output += f", {R_nest_description}"
        full_output = full_output.replace(";", ", ")
        full_output = full_output.replace("is ,", "")
        full_output = process_full_output_text(full_output)
        change_font(doc, full_output)  # Directly add the output to the Word document                         
        if form_paths:
            generate_images_2(doc, image_paths)  # Generate images without y_position
            full_output = "or forms the structure shown below:"
            change_font(doc, full_output)  # Directly add the output to the Word document
            generate_form_images(doc, form_paths)  # Generate form images without y_position
            full_output = forms_description
            change_font(doc, full_output)  # Directly add the output to the Word document
        else:
            generate_images(doc, image_paths)  # Generate images without y_position
    else:
        full_output += f", {R_nest_description}"
        full_output = full_output.replace("is ,", "")
        if form_paths:
            full_output = custom_replace(full_output)
            full_output += f"or forms the structure shown below:"
            full_output = full_output.replace(", or", " or")
            change_font(doc, full_output)  # Directly add the output to the Word document
            generate_form_images(doc, form_paths)  # Generate form images without y_position
            full_output = forms_description
            change_font(doc, full_output)  # Directly add the output to the Word document
        else:
            change_font(doc, full_output)  # Directly add the output to the Word document

    
    if not nest_folder_path or not os.path.isdir(nest_folder_path):
        return y_position

    nested_files = os.listdir(nest_folder_path)
    nested_r_folders = [
        f for f in nested_files if f.startswith("R") and os.path.isdir(os.path.join(nest_folder_path, f))
    ]
    
    processed_folders = {}

    for folder in nested_r_folders:
        if "_" in folder:
            parts = folder.split("_")  # Split the folder name by "_"
        else:
            parts = [folder]  # If no "_" is in the folder name, just use the whole folder name

        for part in parts:
            # If the part starts with "R" and hasn't been processed yet, we add it to processed_folders
            if part.startswith("R") and part not in processed_folders:
                processed_folders[part] = folder
            # If the part doesn't start with "R", it's added to processed_folders regardless
            elif not part.startswith("R"):
                processed_folders[part] = folder

    for i, (number, nested_r_folder) in enumerate(sorted(processed_folders.items())):
        nested_r_folder_path = os.path.join(nest_folder_path, nested_r_folder)

        if os.path.isdir(nested_r_folder_path):
            y_position = process_r_folder_recursively(doc, nested_r_folder, nested_r_folder_path, y_position, index=i)
    
    return y_position



def process_full_output_text(full_output):
    """Format the delimiters in the full_output content."""
    # If full_output contains "Rn is" (where n is any number), do not add "or"
    if re.search(r'\bR\d+\s+is\b', full_output):
        # If "Rn is" is found, replace all ";" and " or" with ","
        return custom_replace(full_output)
    full_output = custom_replace(full_output)
    
    # Ensure "or" is only added when there are multiple items
    if "," in full_output:
        parts = full_output.rsplit(",", 1)  # Split once from the right
        full_output = f"{parts[0]} or {parts[1]}"
    
    return full_output


def find_matching_nest_folder(base_folder_path, part_prefix):
    """
    Searches for folders in `base_folder_path` that contain `part_prefix` and 'nest' in any order.
    Matches folders like `R2_R3_R1_nest`, `R3_R1_nest`, etc., as long as they include the `part_prefix` and 'nest'.
    Returns the path to the nest folder if found, otherwise None.
    """
    try:
        # List subdirectories in the base folder
        subfolders = [f for f in os.listdir(base_folder_path) if os.path.isdir(os.path.join(base_folder_path, f))]

        # Define a regex pattern to match folders that contain the part_prefix and end in '_nest'
        # This pattern checks that the folder name includes part_prefix anywhere before '_nest'
        pattern = re.compile(rf".*{part_prefix}.*_nest$", re.IGNORECASE)

        # Search for the first folder that matches the pattern
        for folder in subfolders:
            if pattern.match(folder):
                return os.path.join(base_folder_path, folder)
    except FileNotFoundError:
        print(f"[Error] The base folder '{base_folder_path}' does not exist.")
    return None


def add_folder_description(doc, part, folder_path, y_position, is_last_folder):
    """Add the folder description to the canvas at the specified y_position."""
    combined_description = ""
    forms_description = ""

    # Initialize a list to store X or Z folder descriptions temporarily
    if not hasattr(add_folder_description, "all_folder_descriptions"):
        add_folder_description.all_folder_descriptions = []

    if "R" in part:
        molecule_id_parts = part.split('_')  
        for i in range(len(molecule_id_parts)):
            current_part = '_'.join(molecule_id_parts[:i+1]) 
            
            # Check if "nest" exists in the folder
            contains_nest_file = any("nest" in file_name for file_name in os.listdir(folder_path))

            combined_description, forms_description = process_rule_description(current_part, folder_path)
            full_output = f"{current_part} is {combined_description}"
            if combined_description == "":
                full_output += f"{generate_description()}"

            if not contains_nest_file:
                if forms_description:  
                    full_output += f" or forms {forms_description}" 
                else:
                    full_output += f"; "  
                change_font(doc, full_output)  

            nest_folder_path = find_matching_nest_folder(folder_path, current_part)
            if nest_folder_path is None:
                continue
            else:
                full_output += f", "  
                _, nest_folder_path, success, R_nest_description, image_paths, form_paths = process_nest_folder(
                    doc, current_part, folder_path, None  
                )              

                if image_paths:
                    full_output += R_nest_description
                    full_output = full_output.replace(";", ", ")
                    full_output = full_output.replace("is ,", "")
                    full_output = process_full_output_text(full_output)
                    change_font(doc, full_output)  
                    
                    if form_paths:
                        generate_images_2(doc, image_paths)  # Generate images for Word document
                        change_font(doc, "or forms the structure shown below:")  # Add text to Word document
                        generate_form_images(doc, form_paths)  # Generate form images for Word document
                        change_font(doc, forms_description)  # Add forms description to Word document
                    else:
                        generate_images(doc, image_paths)  # Generate images for Word document
                else:
                    full_output += R_nest_description
                    full_output = full_output.replace("is ,", "")
                    if form_paths:
                        full_output = custom_replace(full_output)
                        full_output += f"or forms the structure shown below:"
                        full_output = full_output.replace(", or", " or")
                        change_font(doc, full_output)  # Add output to Word document
                        generate_form_images(doc, form_paths)  # Generate form images for Word document
                        change_font(doc, forms_description)  # Add forms description to Word document
                    else:
                        change_font(doc, full_output)  # Add output to Word document
                                                
                
                if os.path.isdir(nest_folder_path):
                    try:
                        nested_files = os.listdir(nest_folder_path)
                        r_folders = sorted(
                            [f for f in nested_files if f.startswith("R") and os.path.isdir(os.path.join(nest_folder_path, f))],
                            key=lambda x: int(subscript_to_int(x[1:])) if subscript_to_int(x[1:]).isdigit() else float('inf')
                        )

                        for index, r_folder in enumerate(r_folders):
                            r_folder_path = os.path.join(nest_folder_path, r_folder)
                            y_position = process_r_folder_recursively(doc, r_folder, r_folder_path, y_position, index)
                    except FileNotFoundError:
                        print(f"[Error] The nest folder '{nest_folder_path}' could not be found.")

    if "X" in part or "Z" in part:
        replacements_path = os.path.join(folder_path, "replacements.txt")
        if os.path.exists(replacements_path):
            with open(replacements_path, "r") as file:
                replacement_lines = [line.replace("*", "").strip() for line in file if line.strip()]
                replacement_lines = list(set(replacement_lines))
                replacement_lines = [line for line in replacement_lines if line]

                common_elements = {"C", "H", "O", "N", "S"}
                halogens = {"F", "Cl", "Br", "I"}
                special_atoms = [line for line in replacement_lines if line not in common_elements and line not in halogens]

                if special_atoms:
                    combined_description = ", ".join(special_atoms)
                else:
                    if "N" in replacement_lines or "C" in replacement_lines and len(replacement_lines) == 1:
                        combined_description = "C or N"
                    else:
                        has_common_elements = any(element in replacement_lines for element in common_elements)
                        has_halogens = any(element in replacement_lines for element in halogens)

                        combined_description = ""
                        if has_common_elements:
                            combined_description += "H, C, N, S or O" if "X" in part else "C, N, S or O"
                        elif has_halogens:
                            combined_description += "a halogen atom"
                        elif len(replacement_lines) > 1:
                            combined_description += ", ".join(replacement_lines[:-1]) + f" or {replacement_lines[-1]}"
                        elif replacement_lines:
                            combined_description += replacement_lines[0]

                if combined_description:
                    add_folder_description.all_folder_descriptions.append((part, combined_description))

    if is_last_folder and hasattr(add_folder_description, "all_folder_descriptions"):
        folder_descriptions = add_folder_description.all_folder_descriptions
        grouped_folders = {}

        for part, description in folder_descriptions:
            grouped_folders.setdefault(description, []).append(part)

        for description, parts in grouped_folders.items():
            if len(parts) > 1:
                part_list = ", ".join(parts[:-1]) + " and " + parts[-1]
                final_description = f"{part_list} are the same or different and each independently represents {description}"
            else:
                final_description = f"{parts[0]} represents {description}"

            # Add ";" if it's not the last description, otherwise add "."
            final_description += ";" if description != list(grouped_folders.keys())[-1] else "."

            # Directly output text to the document
            change_font(doc, final_description)
            
        add_folder_description.all_folder_descriptions = []

    return y_position


def create_claims(output_folder):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"Claims_{timestamp}.docx"
    
    docx_path = os.path.join(output_folder, docx_filename)

    create_word_document(docx_path, output_folder)

    return docx_path
    


